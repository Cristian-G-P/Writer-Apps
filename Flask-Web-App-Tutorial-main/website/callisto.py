import docx
from docx import Document
from .callistoAuxFunctions import *
from datetime import datetime


def formatText(file, dataFromHtml):
    document = Document(PATH_TEMPLATE + "CallistoTemplate.docx")
    docToRead = docx.Document(file)
    setSectionFormat(document, dataFromHtml, docToRead)
    paragraphReadChapterCount = 0
    indentFirstLine = False
    isPreviousAPart = False
    chapterDescription = False
    previousTitle = ''

    if dataFromHtml.ChapterDescription == 'Yes':
        hasChapterDescription = True
    else:
        hasChapterDescription = False

    if dataFromHtml.HalfTitlePage == 'Yes':
        setHalfTitleFormat(document, dataFromHtml)

    if dataFromHtml.TitlePage == 'Yes':
        setTitleFormat(document, dataFromHtml)

    if dataFromHtml.CopyrightPage =='Yes':
        setCopyrightFormat(document, dataFromHtml)

    paragraphsCount = len(docToRead.paragraphs)

    while paragraphReadChapterCount < paragraphsCount - 1:

        if paragraphReadChapterCount == 0:
            paragraphRead = docToRead.paragraphs[paragraphReadChapterCount].text.strip()
            if len(paragraphRead) == 0:
                paragraphReadChapterCount = paragraphReadChapterCount + 1
                continue
        else:
            paragraphRead = paragraphReadNext

        paragraphReadChapterCount = paragraphReadChapterCount + 1

        while paragraphReadChapterCount < paragraphsCount:
            paragraphReadNext = docToRead.paragraphs[paragraphReadChapterCount].text.strip()
            if len(paragraphReadNext) == 0:
                paragraphReadChapterCount = paragraphReadChapterCount + 1
                continue
            break

        paragraghType = setParagraphType(paragraphRead)

        if paragraghType == PART_TEXT:
            setPartFormat(document, paragraphRead)
            isPreviousAPart = True

        elif paragraghType == CHAPTER_TEXT:
            setChapterTitleFormat(document, paragraphRead, isPreviousAPart)
            indentFirstLine = False
            isPreviousAPart = False
            previousTitle = CHAPTER_TEXT
            if hasChapterDescription:
                chapterDescription = True

        elif paragraghType == EXTRA_SECTION :
            setChapterTitleFormat(document, paragraphRead, isPreviousAPart)
            indentFirstLine = False
            isPreviousAPart = False
            chapterDescription = False
            previousTitle = EXTRA_SECTION

        elif paragraghType == SCENE_BREAK_1 or paragraghType == SCENE_BREAK_2:
            setSceneBreakFormat(document, paragraphRead)
            indentFirstLine = False
        else:
            setParagraphFormat(document, paragraphRead, indentFirstLine, paragraphReadNext, chapterDescription, previousTitle, dataFromHtml)
            if chapterDescription:
                indentFirstLine = False
            else:
                indentFirstLine = True
            isPreviousAPart = False
            chapterDescription = False

    document.save(PATH_TEMPORAL + dataFromHtml.bookTitle + FORMATTED_TEXT + '.docx')

    saveCallistoLog(dataFromHtml.bookTitle, dataFromHtml.author)

def saveCallistoLog(book, author):
    now = datetime.now()
    date_time = now.strftime("%m/%d/%Y, %H:%M:%S")
    doc = docx.Document(PATH_LOG + 'CallistoLogs.docx')
    paragraph = date_time + ' --- ' + book + ' --- ' + author
    doc.add_paragraph(paragraph)
    doc.save(PATH_LOG + 'CallistoLogs.docx')
