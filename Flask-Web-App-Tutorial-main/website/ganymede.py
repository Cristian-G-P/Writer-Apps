import re
from datetime import datetime
import docx
from .values import *


class RemoveSpecialCharsValue:
    cleanString = ''
    isWordAccepted = False

class Results:
    paragraphSizeDic = dict()
    paragraphTextDic = dict()
    numberOfWordsOrdered = []
    adverbsCount = dict()
    wordsInChapter = []
    wordsInBook = int
    filterWords = dict()
    glueWords = dict()
    redundantPhrases = dict()
    fillerPhrases = dict()
    nominalizationWords = dict()


def analyzeGrammar(file, dataFromHtml):
    results = Results()
    doc = docx.Document(file)
    wordsCountInChapter = 0
    wordsCountInBook = 0
    previousWord = ''
    adverbsCount = dict()
    paragraphReadChapterCount = 0
    paragraphWords = dict()
    wordsInChapter = dict()
    countChapter = False

    while paragraphReadChapterCount < len(doc.paragraphs):
        paragraphRead = doc.paragraphs[paragraphReadChapterCount].text

        paragraphReadChapterCount = paragraphReadChapterCount + 1
        aux = paragraphRead.strip()

        keys = aux.split(" ")
        wordsCountInBook = wordsCountInBook + len(keys)

        if keys[0] == 'CHAPTER':
            wordsCountInChapter = 0
            countChapter = True
            chapterName = ''
            i = 0
            while i < len(keys):
                chapterName = chapterName + ' ' + keys[i]
                i = i + 1
            wordsInChapter[chapterName] = 0

        results.paragraphSizeDic[paragraphReadChapterCount] = len(keys)

        if len(keys) > PARAGRAPH_NORMAL_SIZE:
            results.paragraphTextDic[paragraphReadChapterCount] = paragraphRead[0:40]

        paragraphReadSplit = paragraphRead.split()
        paragraphReadSplitUpper = []
        ##To put in order the dictionary of words in paragraphs
        str3 = []

        ##iterates every word in the current paragraph
        for i in paragraphReadSplit:
            paragraphReadSplitUpper.append(i.upper())

#################################################################################
            ##Remove special characters
            regex = re.compile('[@_!#$%^&*()<>?/\|}{~:,.]')
            # for i in range(0, len(paragraphReadSplitUpper)):

            if (regex.search(i) == None):
                isWordAccepted = True
                key = i.upper()
            else:
                removeSpecialCharsValue = removeSpecialChars(i.upper())
                key = removeSpecialCharsValue.cleanString

#################################################################################
            ##Chapter counter
            if countChapter:
                wordsCountInChapter = wordsCountInChapter + 1

#################################################################################
            ##Word repetition counter
            if key in paragraphWords.keys():
                paragraphWords[key] = paragraphWords[key] + 1
            else:
                paragraphWords[key] = 1
################################################################################
            # GLUE WORDS
            if key in GLUE_WORDS:
                if key in results.glueWords.keys():
                    results.glueWords[key] = results.glueWords[key] + 1
                else:
                    results.glueWords[key] = 1

#################################################################################
                # FILTER WORDS
            if key in FILTER_WORDS:
                if key in results.filterWords.keys():
                    results.filterWords[key] = results.filterWords[key] + 1
                else:
                    results.filterWords[key] = 1

###################################################################################
            # Calculate number of adverbs
            position = (len(key) - 2)
            if key[position:] == 'LY':

                if key in adverbsCount.keys():
                    adverbsCount[key] = adverbsCount[key] + 1
                else:
                    adverbsCount[key] = 1

###################################################################################
            # REDUNDANT PHRASES
            redundantPhrase = previousWord + ' ' + key
            if redundantPhrase in REDUNDANT_PHRASES:
                if key in results.redundantPhrases.keys():
                    results.redundantPhrases[redundantPhrase] = results.redundantPhrases[redundantPhrase] + 1
                else:
                    results.redundantPhrases[redundantPhrase] = 1
###################################################################################
            # FILLER PHRASES
            fillerPhrase = previousWord + ' ' + key
            if fillerPhrase in FILLER_PHRASES:

                if key in results.fillerPhrases.keys():
                    results.fillerPhrases[fillerPhrase] = results.fillerPhrases[fillerPhrase] + 1
                else:
                    results.fillerPhrases[fillerPhrase] = 1
###################################################################################
            # NOMINALIZATION WORDS
            if key in NOMINALIZATION_WORDS:
                if key in results.nominalizationWords.keys():
                    results.nominalizationWords[key] = results.nominalizationWords[key] + 1
                else:
                    results.nominalizationWords[key] = 1
###################################################################################
            previousWord = key
        if countChapter:
            wordsInChapter[chapterName] = wordsInChapter[chapterName] + wordsCountInChapter
        wordsCountInChapter = 0
###################################################################################

    results.paragraphWords = paragraphWords
    results.numberOfWordsOrdered = str3
    results.adverbsCount = adverbsCount
    results.wordsInChapter = wordsInChapter
    results.wordsInBook = wordsCountInBook

    calculateResults(results, dataFromHtml)


def removeSpecialChars(stringToRemoveChars):
    response = RemoveSpecialCharsValue()

    regex = re.compile('[@_!#$%^&*()<>?/\|}{~:,.’]')

    if (regex.search(stringToRemoveChars) == None):
        q = 0
    else:
        aux1 = stringToRemoveChars.replace('.', '')
        aux2 = aux1.replace(',', '')
        aux3 = aux2.replace(';', '')
        aux4 = aux3.replace('’', '')
        aux5 = aux4.replace('‘', '')
        aux6 = aux5.replace(':', '')
        aux7 = aux6.replace('?', '')
        aux8 = aux7.replace('_', '')
        aux9 = aux8.replace('!', '')
        aux10 = aux9.replace('(', '')
        aux11 = aux10.replace(')', '')
        aux12 = aux11.replace('"', '')
        aux13 = aux12.replace('-', '')
        aux14 = aux13.replace('“', '')
        aux15 = aux14.replace('“', '')

    response.cleanString = aux15

    response.isWordAccepted = True

    return response


def calculateResults(results, dataFromHtml):
    for x in results.paragraphWords.keys():
        results.numberOfWordsOrdered.append(results.paragraphWords[x])

#########################################OVER USED WORDS##########')
    for z in results.paragraphWords.keys():
        results.numberOfWordsOrdered.append(results.paragraphWords[z])

    # results.numberOfWordsOrdered.sort(reverse=True);
    # list out keys and values separately
    key_list = list(results.paragraphWords.keys())
    val_list = list(results.paragraphWords.values())

    val_list2 = list(results.paragraphWords.values())
    val_list2.sort(reverse=True)

    mostUsedWordsDic = dict()

    wordsCounter = 0
    # for i in results.numberOfWordsOrdered:
    ##Si dos palabras tiene la misma cantidad de repeticiones, muestra una sola
    for i in val_list2:
        position = val_list.index(i)
        mostUsedWordsDic[key_list[position]] = val_list[position]

        wordsCounter = wordsCounter + 1
        if wordsCounter > OVER_USED_WORDS_NUM_DISPLAY:
            ##ESTE VALOR DEBERIA SER CONSTANTE
            break

    results.mostUsedWordsDic = mostUsedWordsDic
##############################################################
    printDoc(results, dataFromHtml)

def printDoc(results, dataFromHtml):

    document = docx.Document(PATH_TEMPLATE  + 'GanymedeTemplate.docx')
    document.add_section()

    paragraphWrite = document.add_paragraph(' ')
    paragraphWrite.style = document.styles['BTitle']
    paragraphWrite = document.add_paragraph(' ')
    paragraphWrite.style = document.styles['BTitle']
    paragraphWrite = document.add_paragraph(' ')
    paragraphWrite.style = document.styles['BTitle']
    paragraphWrite = document.add_paragraph(dataFromHtml.bookTitle)
    paragraphWrite.style = document.styles['BTitle']

##############################################################################

    PrintTitleAndDescription(document, "FILLER WORDS", FILLER_WORDS_TEXT)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 1'
    row = table.rows[0].cells
    row[0].text = 'FILLER WORDS'
    row[1].text = 'COUNT'

    for x in FILLER_WORDS:
        if x in results.paragraphWords.keys():
            row = table.add_row().cells
            row[0].text = str(x)
            row[1].text = str(results.paragraphWords[x])

    printLinksTitle(document)

    printLink(document, 'https://www.98thpercentile.com/blog/what-are-verbal-fillers-how-do-they-affect-your-speech/',
              '98thPercentile.com')

    printLink(document, 'https://crowwriter.com/cutting-filler-words-in-writing/', 'crowwriter.com')

    printLink(document, 'https://becomeawritertoday.com/filler-words-list/', 'becomeawritertoday.com')

    printLink(document, 'https://blog.wordvice.com/avoid-fillers-powerful-writing/', 'blog.wordadvice.com')

    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')

##################################################################
    PrintTitleAndDescription(document, "FILTER WORDS", FILTER_WORDS_TEXT)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 1'
    row = table.rows[0].cells
    row[0].text = 'FILTER WORDS'
    row[1].text = 'COUNT'

    for x in results.filterWords:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(x)
        row[1].text = str(results.filterWords[x])

        table.style = 'Medium Grid 3 Accent 1'

    printLinksTitle(document)

    printLink(document, 'https://michaeljmcdonagh.wordpress.com/2014/01/28/writer-unfiltered/',
              'michaeljmcdonagh.wordpress.com')
    printLink(document, 'https://www.masterclass.com/articles/how-to-avoid-unnecessary-filter-words-in-your-writing',
              'masterclass.com')
    printLink(document, 'https://www.rabbitwitharedpen.com/blog/filter-words-in-fiction', 'rabbitwitharedpen.com')
    printLink(document,
              'https://heebel.com/2017/05/28/a-few-easy-strategies-to-remove-those-pesky-filter-words-that-fog-up-your-writing/',
              'heebel.com')

##################################################################
    PrintTitleAndDescription(document, "TAG WORDS", TAG_WORDS_TEXT)
    paragraphWrite = document.add_paragraph(TAG_WORDS_TEXT_2)
    paragraphWrite.style = document.styles['NormalWoSpacing']

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 1'
    row = table.rows[0].cells
    row[0].text = 'TAG WORDS'
    row[1].text = 'COUNT'

    for x in TAG_WORDS:
        if x in results.paragraphWords.keys():
            paragraph = x

            row = table.add_row().cells
            row[0].text = str(x)
            row[1].text = str(results.paragraphWords[x])

    printLinksTitle(document)

    printLink(document, 'https://www.scribophile.com/academy/he-said-she-said-dialog-tags-and-using-them-effectively',
              'scribophile.com')
    printLink(document, 'https://www.michellereneemiller.com/dialogue-tags/', 'michellereneemiller.com')
    printLink(document, 'https://www.nownovel.com/blog/dialogue-words-other-words-for-said/', 'nownovel.com')
    printLink(document, 'https://thewritepractice.com/dialogue-tags/', 'thewritepractice.com')
    printLink(document, 'https://www.wattpad.com/948375492-data-for-you-dear-tips-0-10', 'wattpad.com')

##################################################################
    PrintTitleAndDescription(document, "GLUE WORDS", GLUE_WORDS_TEXT)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 1'
    row = table.rows[0].cells
    row[0].text = 'GLUE WORDS'
    row[1].text = 'COUNT'

    for x in results.glueWords:
        row = table.add_row().cells
        row[0].text = str(x)
        row[1].text = str(results.glueWords[x])

    printLinksTitle(document)

    printLink(document, 'https://readable.com/blog/what-are-glue-words-and-how-do-they-affect-readability/',
              'readable.com')

##################################################################
    PrintTitleAndDescription(document, "NOMINALIZATION WORDS", NOMINALIZATION_WORDS_TEXT)

    printWithDescription(document, 'NOMINALIZATION WORD', 'COUNT', 'DESCRIPTION', results.nominalizationWords,
                         NOMINALIZATION_WORDS)

    printLinksTitle(document)

    printLink(document, 'https://www.thoughtco.com/nominalization-in-grammar-1691430', 'thoughtco.com')
    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')
##################################################################
    PrintTitleAndDescription(document, "ADVERBS ENDING IN LY", ADVERBS_TEXT)

    table = document.add_table(rows=1, cols=2)

    row = table.rows[0].cells
    row[0].text = 'ADVERBS (-LY)'
    row[1].text = 'COUNT'

    for x in results.adverbsCount:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(x)
        row[1].text = str(results.adverbsCount[x])

        table.style = 'Medium Grid 3 Accent 1'

    printLinksTitle(document)

    printLink(document, 'https://glcubel.com/2020/05/12/a-guide-to-removing-adverbs/', 'glcubel.com')
##################################################################
    PrintTitleAndDescription(document, "OVERUSED WORDS (PART I)", OVERUSED_WORDS_TEXT)

    table = document.add_table(rows=1, cols=2)

    row = table.rows[0].cells
    row[0].text = 'WORD'
    row[1].text = 'COUNT'

    for x in results.mostUsedWordsDic:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(x)
        row[1].text = str(results.mostUsedWordsDic[x])

        table.style = 'Medium Grid 3 Accent 1'
##################################################################
    PrintTitleAndDescription(document, "OVERUSED WORDS (PART II)", OVERUSED_WORDS_TEXT_2)

    table = document.add_table(rows=1, cols=2)

    row = table.rows[0].cells
    row[0].text = 'WORD'
    row[1].text = 'COUNT'

    for x in results.mostUsedWordsDic:
        if x not in COMMON_WORDS:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(x)
            row[1].text = str(results.mostUsedWordsDic[x])

    table.style = 'Medium Grid 3 Accent 1'

##################################################################
    PrintTitleAndDescription(document, "UNIQUE WORDS", UNIQUE_WORDS_TEXT)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 1'

    row = table.rows[0].cells
    row[0].text = 'NUMBER OF UNIQUE WORDS'
    row[1].text = str(len(results.numberOfWordsOrdered))

    printLinksTitle(document)

    printLink(document,
              'https://www.lingholic.com/how-many-words-do-i-need-to-know-the-955-rule-in-language-learning-part-2/',
              'lingholic.com')
##################################################################
    PrintTitleAndDescription(document, "FILLER PHRASES", FILLER_PHRASES_TEXT)

    printWithDescription(document, 'FILLER PHRASE', 'COUNT', 'DESCRIPTION', results.fillerPhrases, FILLER_PHRASES)

    printLinksTitle(document)

    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')
##################################################################
    PrintTitleAndDescription(document, "REDUNDANT PHRASES", REDUNDANT_PHRASES_TEXT)

    printWithDescription(document, 'REDUNDANT PHRASE', 'COUNT', 'DESCRIPTION', results.redundantPhrases,
                         REDUNDANT_PHRASES)

    printLinksTitle(document)

    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')
##################################################################
    PrintTitleAndDescription(document, "PARAGRAPH LENGTH", PARAGRAPH_SIZE_TEXT)
    document.add_paragraph('')

    table = document.add_table(rows=1, cols=3)
    table.style = 'Medium Grid 3 Accent 1'

    row = table.rows[0].cells
    row[0].text = 'PARAGRAPH'
    row[1].text = 'WORD COUNT'
    row[2].text = 'STARTS WITH'

    for x in results.paragraphSizeDic:
        if (results.paragraphSizeDic[x] > PARAGRAPH_NORMAL_SIZE):
            row = table.add_row().cells

            row[0].text = str(x)
            row[1].text = str(results.paragraphSizeDic[x])
            row[2].text = str(results.paragraphTextDic[x])

    printLinksTitle(document)

    printLink(document, 'https://www.masterclass.com/articles/how-long-is-a-paragraph-explained', 'masterclass.com')
###########################################################################
    PrintTitleAndDescription(document, "CHAPTER LENGTH", CHAPTER_SIZE_TEXT)

    table = document.add_table(rows=1, cols=3)

    row = table.rows[0].cells
    row[0].text = 'CHAPTER'
    row[1].text = 'WORD COUNT'
    row[2].text = 'COMMENTS'

    for x in results.wordsInChapter:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(x)
        row[1].text = str(results.wordsInChapter[x])
        if (results.wordsInChapter[x] > CHAPTER_NORMAL_SIZE):
            row[2].text = 'Exceeded'
        else:
            row[2].text = 'OK'

        table.style = 'Medium Grid 3 Accent 1'

    printLinksTitle(document)

    printLink(document, 'https://blog.reedsy.com/how-long-should-a-chapter-be/', 'blog.reedsy.com')
########################################################################
    PrintTitleAndDescription(document, "BOOK LENGTH", BOOK_LENGTH_TEXT)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 1'
    row = table.rows[0].cells
    row[0].text = 'YOUR BOOK WORD COUNT'
    row[1].text = str(results.wordsInBook)
    document.add_paragraph('')

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 1'

    row = table.rows[0].cells
    row[0].text = 'BOOK TYPE'
    row[1].text = 'WORDS'

    table.add_row().cells
    row = table.rows[1].cells
    row[0].text = 'FLASH FICTION'
    row[1].text = '300 - 1500'

    table.add_row().cells
    row = table.rows[2].cells
    row[0].text = 'SHORT STORY'
    row[1].text = '1500 - 30000'

    table.add_row().cells
    row = table.rows[3].cells
    row[0].text = 'NOVELLA'
    row[1].text = '30000 - 50000'

    table.add_row().cells
    row = table.rows[4].cells
    row[0].text = 'NOVEL'
    row[1].text = '50000 - 110000'

    table.add_row().cells
    row = table.rows[5].cells
    row[0].text = 'EPIC STORY'
    row[1].text = '110000 - 800000'

    printLinksTitle(document)

    printLink(document, 'https://prowritingaid.com/art/1243/genre-book-length-.aspx', 'prowritingaid.com')

##################################################################
    PrintTitleAndDescription(document, "CLUNKY CONSTRUCTIONS", CLUNKY_CONSTRUCTION_TEXT)

    printSectionWithoutCases(document, 'CLUNKY CONSTRUCTION', CLUNKY_CONSTRUCTION_TEXT, 'CLUNKY CONSTRUCTION',
                             'DESCRIPTION', CLUNKY_CONSTRUCTIONS)

    printLinksTitle(document)

    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')

##################################################################
    PrintTitleAndDescription(document, "EMPTY PHRASES", EMPTY_PHRASES_TEXT)

    printSectionWithoutCases(document, 'EMPTY PHRASES', EMPTY_PHRASES_TEXT, 'EMPTY PHRASE', 'DESCRIPTION',
                             EMPTY_PHRASES)

    printLinksTitle(document)

    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')
##################################################################
    PrintTitleAndDescription(document, "NEGATIVE CONSTRUCTIONS", NEGATIVE_CONSTRUCTION_TEXT)

    printSectionWithoutCases(document, 'NEGATIVE CONSTRUCTIONS', NEGATIVE_CONSTRUCTION_TEXT, 'NEGATIVE CONSTRUCTION',
                             'DESCRIPTION', NEGATIVE_CONSTRUCTIONS)

    printLinksTitle(document)

    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')
##################################################################
    PrintTitleAndDescription(document, "MORE FILLER PHRASES", FILLER_PHRASES_TEXT_2)

    printSectionWithoutCases(document, 'FILLER PHRASES - PART II', FILLER_PHRASES_TEXT_2, 'FILLER PHRASE',
                             'DESCRIPTION', FILLER_PHRASES2)

##################################################################
    PrintTitleAndDescription(document, "SENTENCE LENGTH", SENTENCE_LENGTH_TEXT)

    printLinksTitle(document)

    printLink(document, 'https://www.thoughtco.com/sentence-length-grammar-and-composition-1691948', 'thoughtco.com')
##################################################################
    PrintTitleAndDescription(document, "PASSIVE VOICE", PASSIVE_VOICE_TEXT)

    printLinksTitle(document)

    printLink(document, 'https://yoast.com/the-passive-voice-what-is-it-and-how-to-avoid-it/', 'yoast.com')

##################################################################
    PrintTitleAndDescription(document, "SHOW DON'T TELL", SHOW_DONT_TELL_TEXT)

    printLinksTitle(document)

    printLink(document, 'https://jerryjenkins.com/show-dont-tell/', 'jerryjenkins.com')
    printLink(document, 'https://blog.reedsy.com/show-dont-tell/', 'blog.reedsy.com')
    printLink(document, 'https://www.nownovel.com/blog/showing-vs-telling/', 'nownovel.com')

 ########################################################################################################
    saveGanymedeLog(dataFromHtml.bookTitle)

    document.save(PATH_TEMPORAL + dataFromHtml.bookTitle + ANALYZED_TEXT + '.docx')

########################################################################################################

def printLink(document, link, text):
    p = document.add_paragraph(style='List Bullet')
    add_hyperlink(p, link, text, '4122ff', True)
    font = p.style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

def PrintTitleAndDescription(document, title, description):
    document.add_page_break()
    document.add_heading(title)
    paragraph = description
    paragraphWrite = document.add_paragraph(paragraph)
    paragraphWrite.style = document.styles['NormalWoSpacing']


def printLinksTitle(document):
    paragraphWrite = document.add_paragraph('')
    paragraph = ('Visit these links for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    paragraphWrite.style = document.styles['NormalWoSpacing']


def printSection(document, paragraph, text, col1, col2, results):
    document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = text
    document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 1'

    row = table.rows[0].cells
    row[0].text = col1
    row[1].text = col2

    for x in results:
        row = table.add_row().cells
        row[0].text = str(x)
        row[1].text = str(results[x])


def printWithDescription(document, col1, col2, col3, results, dic):
    table = document.add_table(rows=1, cols=3)
    table.style = 'Medium Grid 3 Accent 1'

    row = table.rows[0].cells
    row[0].text = col1
    row[1].text = col2
    row[2].text = col3

    for x in results:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(x)
        row[1].text = str(results[x])
        row[2].text = str(dic[x])


def printSectionWithoutCases(document, title, text, col1, col2, dic):
    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 1'

    row = table.rows[0].cells
    row[0].text = col1
    row[1].text = col2

    for x in dic:
        row = table.add_row().cells
        row[0].text = str(x)
        row[1].text = str(dic[x])


############################################################
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    # if not underline:
    # u = docx.oxml.shared.OxmlElement('w:u')
    # u.set(docx.oxml.shared.qn('w:val'), 'none')
    # rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def setTitleFormatGrammar(document, paragraphRead):
    document.add_heading(paragraphRead)

def saveGanymedeLog(book):
    now = datetime.now()
    date_time = now.strftime("%m/%d/%Y, %H:%M:%S")

    doc = docx.Document(PATH_LOG + 'GanymedeLogs.docx')

    paragraph = date_time + ' --- ' + book
    doc.add_paragraph(paragraph)

    doc.save(PATH_LOG + 'GanymedeLogs.docx')
