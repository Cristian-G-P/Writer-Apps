import docx
from docx import Document

from .auxFunctions import *
from .values import *
from datetime import datetime



def formatText(file, dataFromHtml):

    document = Document()
    doc = docx.Document('website/temporal/' + file.filename)

    document.styles
    setSectionFormat(document, dataFromHtml, doc)

    paragraphReadChapterCount = 0
    indentFirstLine = False
    isPreviousAPart = False
    chapterDescription= False

    if dataFromHtml.ChapterDescription == 'Yes':
        hasChapterDescription = True
    else:
        hasChapterDescription = False

    if dataFromHtml.HalfTitlePage =='Yes':
        setHalfTitleFormat(document, dataFromHtml)

    if dataFromHtml.TitlePage =='Yes':
        setTitleFormat(document, dataFromHtml)

    if dataFromHtml.CopyrightPage =='Yes':
        setCopyrightFormat(document, dataFromHtml)

    style = document.styles.add_style("Chapter Description", WD_STYLE_TYPE.PARAGRAPH)
    style.hidden = False
    style.quick_style = True
    style.priority = 3

    chapterStyle = document.styles.add_style("Chapter Title", WD_STYLE_TYPE.PARAGRAPH)
    styles = document.styles
    chapterStyle.base_style = styles['Heading 1']
    chapterStyle.hidden = False
    chapterStyle.quick_style = True
    chapterStyle.priority = 1
    #chapterStyle.font = dataFromHtml.ChapterTitleFont



    #setDedicationFormat(document, titleText, subtitleText)

    while paragraphReadChapterCount < len(doc.paragraphs):
        paragraphRead = doc.paragraphs[paragraphReadChapterCount].text.strip()

        isParagrapahChapterTitle = isParagraphAChapter(paragraphReadChapterCount, doc)

        isParagraphDedication = isParagraphADedication(paragraphReadChapterCount, doc)

        isParagraphPart = isParagraphAPart(paragraphReadChapterCount, doc)
        #print(paragraphRead)
        if isParagraphPart:
            setPartFormat(document, paragraphRead)
            isPreviousAPart = True
         #   print('part')

        elif chapterDescription:
            if (len(paragraphRead.strip()) > 0):
                setChapterDescriptionFormat(document, paragraphRead, isPreviousAPart, dataFromHtml)
                chapterDescription = False
          #      print('descriptio')

        elif isParagrapahChapterTitle:
            setChapterTitleFormat(document, paragraphRead, isPreviousAPart, dataFromHtml)
            indentFirstLine = False
            isPreviousAPart = False
            if hasChapterDescription:
                chapterDescription = True
           # print('chaptretitle')

        elif isParagraphDedication:
            setChapterTitleFormat(document, paragraphRead, isPreviousAPart, dataFromHtml)
            indentFirstLine = False
            isPreviousAPart = False
            chapterDescription = False
            #print('dedication')

        else:
            setChapterFormat(document, paragraphRead, doc, paragraphReadChapterCount, indentFirstLine, dataFromHtml,chapterDescription)
            if(len(paragraphRead.strip()) > 0):
                indentFirstLine = True
                isPreviousAPart = False
                chapterDescription = False
            #print('chapter')

        paragraphReadChapterCount = paragraphReadChapterCount + 1

    #setSectionFormat(document, dataFromHtml, countWords)
    #print(countWords)
    document.save('website/temporal/' + dataFromHtml.bookTitle+ ' - Formatted by Callisto' + '.docx')

    saveCallistoLog(dataFromHtml.bookTitle, dataFromHtml.author)
def saveCallistoLog(book, author):
    now = datetime.now()
    date_time = now.strftime("%m/%d/%Y, %H:%M:%S")


    doc = docx.Document('website/temporal/CallistoLogs.docx')

    paragraph = date_time + ' --- ' + book + ' --- ' + author
    doc.add_paragraph(paragraph)
    doc.save('website/temporal/CallistoLogs.docx')
