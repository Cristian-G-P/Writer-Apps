

from .values import *
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def isParagraphAChapter(paragraphReadChapterCount, doc):
    isParagraphChapterTitle: bool = False

    if (paragraphReadChapterCount) < len(doc.paragraphs):
        paragraphReadAux = doc.paragraphs[paragraphReadChapterCount].text.strip()
        firstOccurenceAux: int = paragraphReadAux.upper().find(CHAPTER_TEXT.upper())

        if firstOccurenceAux == 0:
            isParagraphChapterTitle = True
        else:
            isParagraphChapterTitle = False

    return isParagraphChapterTitle



def isParagraphADedication(paragraphReadChapterCount, doc):
    isParagraphADedication: bool = False

    if (paragraphReadChapterCount) < len(doc.paragraphs):
        paragraphReadAux = doc.paragraphs[paragraphReadChapterCount].text.strip()
        firstOccurenceAux: int = paragraphReadAux.upper().find(DEDICATION_TEXT.upper())

        if firstOccurenceAux == 0:
            isParagraphADedication = True

        else:
            isParagraphADedication = False

    return isParagraphADedication

def isParagraphAPart(paragraphReadChapterCount, doc):
    isParagraphAPart: bool = False

    if (paragraphReadChapterCount) < len(doc.paragraphs):
        paragraphReadAux = doc.paragraphs[paragraphReadChapterCount].text.strip()
        firstOccurenceAux: int = paragraphReadAux.upper().find(PART_TEXT.upper())

        if firstOccurenceAux == 0:
            isParagraphAPart = True

        else:
            isParagraphAPart = False

    return isParagraphAPart

def isPreviousNotEmptyParagraphAPart(paragraphReadChapterCount, doc):
    isParagraphAPart: bool = False

    while (paragraphReadChapterCount > 0):
        paragraphReadAux = doc.paragraphs[paragraphReadChapterCount - 1].text.strip()
        if (len(paragraphReadAux)>0):
            firstOccurenceAux: int = paragraphReadAux.upper().find(PART_TEXT.upper())

            if firstOccurenceAux == 0:
                isParagraphAPart = True
            else:
                isParagraphAPart = False
            break

        paragraphReadChapterCount = paragraphReadChapterCount - 1

    return isParagraphAPart


def setIdentitation(paragraphReadChapterCount, doc, indent):
    indentFirstLine = indent

    previousParagraphNotEmpty = doc.paragraphs[paragraphReadChapterCount - 1].text.strip()

    counter = 1

    while counter < paragraphReadChapterCount:
        if (len(doc.paragraphs[paragraphReadChapterCount - counter].text.strip()) > 0):
            previousParagraphNotEmpty = doc.paragraphs[paragraphReadChapterCount - counter].text
            break
        counter = counter + 1

    if previousParagraphNotEmpty.strip() == '***':
        indentFirstLine = False

    if previousParagraphNotEmpty.strip() == '###':
        indentFirstLine = False

    return indentFirstLine

def setChapterTitleFormatAnalysis(document, paragraphRead):
    paragraphWrite = document.add_heading(paragraphRead)

    paragraph_format = paragraphWrite.paragraph_format
    paragraph_format.space_before = Pt(10)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # paragraphWrite.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
    font = paragraphWrite.style.font
    font.name = font_name
    font.size = Pt(fontSizeTitle)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color

    return paragraph_format




def setHalfTitleFormat(document, dataFromHtml):

    paragraphWrite = document.add_paragraph(dataFromHtml.bookTitle)
    style = document.styles.add_style('HalfTitle1', WD_STYLE_TYPE.PARAGRAPH)
    paragraphWrite.style = style

    paragraph_format = paragraphWrite.paragraph_format
    paragraph_format.space_before = Pt(150)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = paragraphWrite.style.font
    font.name = dataFromHtml.HalfTitleFont
    size = int(dataFromHtml.HalfTitleFontSize)
    font.size = Pt(size)
    font.bold = False
    font.italic = False
    font.underline = False
    font.color.rgb = RGBColor(0, 0, 0)

    document.add_section()
    document.add_section()

def setTitleFormat(document, dataFromHtml):
    paragraphWrite1 = document.add_paragraph(dataFromHtml.bookTitle)
    #paragraph_format = paragraphWrite.paragraph_format
    style = document.styles.add_style('TitlePageTitle', WD_STYLE_TYPE.PARAGRAPH)
    paragraphWrite1.style = style

    paragraph_format = paragraphWrite1.paragraph_format
    paragraph_format.space_before = Pt(150)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    font = paragraphWrite1.style.font
    font.name = dataFromHtml.TitlePageFont
    size = int(dataFromHtml.TitlePageFontSize)
    font.size = Pt(size)
    font.bold = False
    font.italic = False
    font.underline = False
    font.color.rgb = RGBColor(0, 0, 0)

    paragraphWrite2 = document.add_paragraph(dataFromHtml.bookSubTitle)
    #paragraph_format = paragraphWrite2.paragraph_format
    style = document.styles.add_style('TitlePageSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    paragraphWrite2.style = style

    paragraph_format = paragraphWrite2.paragraph_format
    paragraph_format.space_before = Pt(10)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = paragraphWrite2.style.font
    font.name = dataFromHtml.SubtitleFont
    size = int(dataFromHtml.SubtitleFontSize)
    font.size = Pt(size)
    font.bold = False
    font.italic = False
    font.underline = False
    font.color.rgb = RGBColor(0, 0, 0)

    document.add_paragraph('')
    paragraphWrite3 = document.add_paragraph(dataFromHtml.author)
    #paragraph_format = paragraphWrite3.paragraph_format
    style = document.styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
    paragraphWrite3.style = style

    paragraph_format = paragraphWrite3.paragraph_format
    paragraph_format.space_before = Pt(10)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = paragraphWrite3.style.font
    font.name = dataFromHtml.AuthorFont
    size = int(dataFromHtml.AuthorFontSize)
    font.size = Pt(size)
    font.bold = False
    font.italic = False
    font.underline = False
    font.color.rgb = RGBColor(0, 0, 0)

    document.add_section()
    #document.add_section()


def setCopyrightFormat(document, dataFromHtml):
    style = document.styles.add_style('Copyright1', WD_STYLE_TYPE.PARAGRAPH)


    paragraphWrite = document.add_paragraph(dataFromHtml.bookTitle +' by ' + dataFromHtml.author)
    paragraphWrite.style = style
    paragraph_format = paragraphWrite.paragraph_format
    paragraph_format.space_before = Pt(250)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragraphWrite = document.add_paragraph('Copyright 2022 by ' + dataFromHtml.author)
    paragraphWrite.style = style

    paragraphWrite = document.add_paragraph(COPYRIGHT_TEXT)
    paragraphWrite.style = style
    ISBN = str(dataFromHtml.ISBN)

    paragraphWrite = document.add_paragraph('ISBN: ' + ISBN)
    paragraphWrite.style = style

    paragraph_format = paragraphWrite.paragraph_format

    font = paragraphWrite.style.font
    font.name = dataFromHtml.CopyrightFont
    size = int(dataFromHtml.CopyrightFontSize)
    font.size = Pt(size)
    font.bold = False
    font.italic = False
    font.underline = False
    font.color.rgb = RGBColor(0, 0, 0)

    new_section = document.add_section()


def setPartFormat(document, paragraphRead):
    #document.add_section()
    #paragraphWrite = document.add_heading(paragraphRead)
    paragraphWrite = document.add_paragraph(paragraphRead)
    #style = document.styles.add_style('HalfTitle1', WD_STYLE_TYPE.PARAGRAPH)
    #paragraphWrite.style = style

    paragraph_format = paragraphWrite.paragraph_format
    #paragraph_format.space_before = Pt(100)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = paragraphWrite.style.font
    #font.name = dataFromHtml.HalfTitleFont
    #size = int(dataFromHtml.HalfTitleFontSize)
    font.size = Pt(14)
    #font.bold = False
    #font.italic = False
    #font.underline = False
    #font.color.rgb = RGBColor(0, 0, 0)

    #document.add_section()
    #document.add_section()



def setChapterTitleFormat(document, paragraphRead, isPreviousAPart, dataFromHtml):
    paragraphWrite = document.add_paragraph(paragraphRead)
    paragraphWrite.style = document.styles['Chapter Title']

    #paragraphWrite = document.add_heading(paragraphRead)

    paragraph_format = paragraphWrite.paragraph_format
    if isPreviousAPart:
        paragraph_format.space_before = Pt(45)
    else:
        paragraph_format.space_before = Pt(70)

    paragraph_format.space_after = Pt(20)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #paragraphWrite.style = document.styles("CHAPTER_DESCRIPTION")
    font = paragraphWrite.style.font
    font.name = dataFromHtml.ChapterTitleFont
    size = int(dataFromHtml.ChapterTitleFontSize)
    font.size = Pt(size)
    font.bold = False
    font.italic = False
    font.underline = font_underline
    font.color.rgb = color

    #new_section = document.add_section()
    #section = document.sections[0]

    #return paragraph_format

def setChapterDescriptionFormat(document, paragraphRead, isPreviousAPart, dataFromHtml):


    paragraphWrite = document.add_paragraph(paragraphRead)
    paragraphWrite.style = document.styles['Chapter Description']

    paragraph_format = paragraphWrite.paragraph_format

    paragraph_format.space_before = Pt(10)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph_format.right_indent = Inches(0.7)
    paragraph_format.left_indent = Inches(0.7)





    #paragraphWrite.style = document.styles['Intense Quote']
    font = paragraphWrite.style.font
    font.name = dataFromHtml.ChapterDescriptionFont
    size = int(dataFromHtml.ChapterDescriptionFontSize)
    font.size = Pt(size)
    font.bold = False
    font.italic = True
    #font.underline = font_underline
    font.color.rgb = RGBColor(0, 0, 0)


    #new_section = document.add_section()
    #section = document.sections[0]

    #return paragraph_format


def setChapterFormat(document, paragraphRead, doc, paragraphReadChapterCount, indentFirstLine, dataFromHtml, chapterDescription):

    if (len(paragraphRead.strip()) > 0):

        paragraphWrite = document.add_paragraph(paragraphRead)
        font = paragraphWrite.style.font
        font.name = dataFromHtml.ParagraphFont
        size = int(dataFromHtml.ParagraphFontSize)
        font.size = Pt(size)

        font.bold = False
        if dataFromHtml.ChapterDescription == 'Yes' and chapterDescription:
            font.italic = True

        else:
            font.italic = False

        font.underline = False
        font.color.rgb = RGBColor(0, 0, 0)

        paragraph_format = paragraphWrite.paragraph_format
        interlining = float(dataFromHtml.Interlining)
        paragraph_format.line_spacing = interlining

        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next

        if (paragraphRead.strip() == '***') or (paragraphRead.strip() == '###'):
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif ( (dataFromHtml.Justification == 'CENTER') or (dataFromHtml.ChapterDescription == 'Yes' and chapterDescription)):
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif (dataFromHtml.Justification == 'JUSTIFY') :
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        indentFirstLine = setIdentitation(paragraphReadChapterCount, doc, indentFirstLine)




        if indentFirstLine:
            paragraph_format.first_line_indent = Pt(18)
        #else:
         #   paragraph_format.first_line_indent = Pt(1)

        indentFirstLine = True

    isPreviousAPart = False
    isPreviousAPart = isPreviousNotEmptyParagraphAPart(paragraphReadChapterCount, doc)
    isNextParagrapahChapterTitle = isParagraphAChapter(paragraphReadChapterCount + 1, doc)
    isNextParagrapahPart = isParagraphAPart(paragraphReadChapterCount + 1, doc)


    if (isNextParagrapahChapterTitle or isNextParagrapahPart) and (not isPreviousAPart):
        new_section = document.add_section()
        #section = document.sections[0]



def setSectionFormat(document, dataFromHtml, doc):
    section = document.sections[0]
    paragraphReadChapterCount=0
    countWords = 0

    while paragraphReadChapterCount < len(doc.paragraphs):
        paragraphRead = doc.paragraphs[paragraphReadChapterCount].text.strip()
        count = paragraphRead.split(" ")
        countWords = countWords + len(count)
        paragraphReadChapterCount = paragraphReadChapterCount + 1

    if dataFromHtml.PageSize == '5" x 8" (12.7 x 20.32 cm)':
        section.page_width = Inches(5)
        section.page_height = Inches(8)
    elif dataFromHtml.PageSize == '5.25" x 8"(13.34 x 20.32 cm)':
        section.page_width = Inches(5.25)
        section.page_height = Inches(8)
    elif dataFromHtml.PageSize =='5.5" x 8.5"(13.97 x 21.59 cm)':
        section.page_width = Inches(5.5)
        section.page_height = Inches(8.5)
    elif dataFromHtml.PageSize == '6" x 9" (15.24 x 22.86 cm)':
        section.page_width = Inches(6)
        section.page_height = Inches(9)
    elif dataFromHtml.PageSize == '7" x 10" x (17.78 x 25.4 cm)':
        section.page_width = Inches(7)
        section.page_height = Inches(10)
    elif dataFromHtml.PageSize == '8" x 10" (20.32 x 25.4 cm)':
        section.page_width = Inches(8)
        section.page_height = Inches(10)
    elif dataFromHtml.PageSize == 'Letter 8.5" x 11" (21,59 x 27,94 cm)':
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    elif dataFromHtml.PageSize == 'A4 8.27" x 11.69" (21 x 29,7 cm)':
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    if dataFromHtml.OutsideMargin == '0.3 in (7.6 mm)':
        section.right_margin = Inches(0.3)
    if dataFromHtml.OutsideMargin == '0.25 in (6.4 mm)':
        section.right_margin = Inches(0.25)

    pageCount = countWords/250

    if pageCount < 150:
        marginCalc = 0.375
    elif pageCount < 300:
        marginCalc = 0.5
    elif pageCount < 500:
        marginCalc = 0.625
    elif pageCount < 700:
        marginCalc = 0.75
    elif pageCount < 828:
        marginCalc = 0.875
    else:
        marginCalc= 0.875

    margin = 0.375

    if dataFromHtml.InsideMargin =="0.375 in (9.6 mm)":
        margin = 0.375
    if dataFromHtml.InsideMargin =="0.5 in (12.7 mm)":
        margin = 0.5
    if dataFromHtml.InsideMargin =="0.625 in (15.9 mm)":
        margin = 0.625
    if dataFromHtml.InsideMargin =="0.75 in (19.1 mm)":
        margin = 0.75
    if dataFromHtml.InsideMargin =="0.875 in (22.3 mm)":
        margin = 0.875

    if margin > marginCalc:
        section.left_margin = Inches(margin)
    else:
        section.left_margin = Inches(marginCalc)
    #section.left_margin = Inches(left_margin)
    #if dataFromHtml.InsideMargin > 0.375

    #section.right_margin = Inches(right_margin)

    section.different_first_page_header_footer = True
    document.settings.odd_and_even_pages_header_footer = True

    header = section.header
    footer = section.footer
    evenHeader = section.even_page_header
    evenFooter = section.even_page_footer

    header.is_linked_to_previous = False


    textHeader = header.paragraphs[0]
    textHeader.text = dataFromHtml.bookTitle

    textEvenHeader = evenHeader.paragraphs[0]
    textEvenHeader.text = dataFromHtml.author

    textFooter = footer.paragraphs[0]
    textFooter.text = 'Odd Page Number'

    textEvenFooter = evenFooter.paragraphs[0]
    textEvenFooter.text = 'Even Page Number'


#############################################################3

def setTitleFormatGrammar(document, paragraphRead):
##################33
    paragraphWrite = document.add_heading(paragraphRead)

    paragraph_format = paragraphWrite.paragraph_format
    paragraph_format.space_before = Pt(30)
    paragraph_format.space_after = Pt(10)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #paragraphWriteG.style = document.styles.add_style('TitleGrammar', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraphWrite.style.font
    font.name = 'Cambria'
    font.size = Pt(20)
    font.bold = False
    font.italic = False
    font.underline = False
    font.color.rgb = RGBColor(255,255,255)

################



    #paragraphWrite = document.add_heading(paragraphRead)
    #font = paragraphWrite.style.font

    #font.name = font_name

    #paragraph_format = paragraphWrite.paragraph_format
    #paragraph_format.space_before = Pt(10)
    #paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#############33
    # Create XML element
    shd = OxmlElement('w:shd')

    # Add attributes to the element
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), '0,0,0')
    shd.set(qn('w:fill'), '4F81BD')

    paragraphWrite.paragraph_format.element.get_or_add_pPr()
    paragraphWrite.paragraph_format.element.pPr.append(shd)
    ##############

    # paragraphWrite.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)

    #font.size = Pt(fontSizeTitle)
    #font.bold = font_bold
    #font.italic = True
    #font.underline = font_underline
    #font.color.rgb = RGBColor(255,255,255)
    #paragraphWrite.font.rgb = RGBColor(0,0,0)

    return paragraph_format

def setTitleFormatGrammar1(document, paragraphRead):
##################33
    paragraphWrite = document.add_heading(paragraphRead)

    paragraph_format = paragraphWrite.paragraph_format
    paragraph_format.space_before = Pt(5)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraphWrite.style = document.styles.add_style('TitleGrammar', WD_STYLE_TYPE.PARAGRAPH)
    font = paragraphWrite.style.font
    font.name = 'Cambria'
    font.size = Pt(26)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color


    #shd = OxmlElement('w:shd')

# Add attributes to the element
    #shd.set(qn('w:val'), 'clear')
    #shd.set(qn('w:color'), '0,0,0')
    #shd.set(qn('w:fill'), '4F81BD')

    #paragraphWrite.paragraph_format.element.get_or_add_pPr()
    #paragraphWrite.paragraph_format.element.pPr.append(shd)