from .values import *
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def setParagraphType(paragraph):
    if paragraph[:7].upper() == CHAPTER_TEXT:
        type = CHAPTER_TEXT
    elif paragraph[:10].upper() == DEDICATION_TEXT:
        type = EXTRA_SECTION
    elif paragraph[:14].upper() == ACKNOWLEDGMENT_TEXT:
        type = EXTRA_SECTION
    elif paragraph[:9].upper() == BIOGRAPHY_TEXT:
        type = EXTRA_SECTION
    elif paragraph[:16].upper() == ABOUT_THE_AUTHOR_TEXT:
        type = EXTRA_SECTION
    elif paragraph[:4].upper() == PART_TEXT:
        type = PART_TEXT
    elif paragraph[:3] == SCENE_BREAK_1:
        type = SCENE_BREAK_1
    elif paragraph[:3] == SCENE_BREAK_2:
        type = SCENE_BREAK_2
    else:
        type = PARAGRAPH

    return type


def setHalfTitleFormat(document, dataFromHtml):
    paragraphWrite = document.add_paragraph(dataFromHtml.bookTitle)
    paragraphWrite.style = document.styles['Book Title']

    document.add_section()
    document.add_section()

def setTitleFormat(document, dataFromHtml):
    paragraphWrite = document.add_paragraph(dataFromHtml.bookTitle)
    paragraphWrite.style = document.styles['Book Title']

    paragraphWrite2 = document.add_paragraph(dataFromHtml.bookSubTitle)
    paragraphWrite2.style = document.styles['Book Subtitle']

    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')

    paragraphWrite3 = document.add_paragraph(dataFromHtml.author.upper())
    paragraphWrite3.style = document.styles['Author']

    document.add_section()

def setCopyrightFormat(document, dataFromHtml):
    i= 0
    while i <= 22:
        document.add_paragraph('')
        i = i + 1

    paragraphWrite = document.add_paragraph(dataFromHtml.bookTitle +' by ' + dataFromHtml.author)
    paragraphWrite.style = document.styles['Copyright']

    paragraphWrite = document.add_paragraph('Copyright© 2022 ' + dataFromHtml.author)
    paragraphWrite.style = document.styles['Copyright']

    paragraphWrite = document.add_paragraph(COPYRIGHT_TEXT)
    paragraphWrite.style = document.styles['Copyright']

    paragraphWrite = document.add_paragraph('ISBN: XXXX-XXXX')
    paragraphWrite.style = document.styles['Copyright']

    document.add_section()


def setPartFormat(document, paragraphRead):
    paragraphWrite = document.add_paragraph(paragraphRead)
    paragraphWrite.style = document.styles['Part']

def setChapterTitleFormat(document, paragraphRead, isPreviousAPart):
    paragraphWrite = document.add_paragraph(paragraphRead)
    paragraphWrite.style = document.styles['Chapter Title']

    paragraph_format = paragraphWrite.paragraph_format
    if isPreviousAPart:
        paragraph_format.space_before = Pt(45)
    else:
        paragraph_format.space_before = Pt(70)

    paragraph_format.space_after = Pt(20)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def setChapterDescriptionFormat(document, paragraphRead):
    paragraphWrite = document.add_paragraph(paragraphRead)
    paragraphWrite.style = document.styles['Chapter Description']

    paragraph_format = paragraphWrite.paragraph_format

    paragraph_format.space_before = Pt(10)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph_format.right_indent = Inches(0.7)
    paragraph_format.left_indent = Inches(0.7)

def setParagraphFormat(document, paragraphRead, indentFirstLine, paragrapghReadNext, chapterDescription, previousTitle, data):
    paragraphWrite = document.add_paragraph(paragraphRead)

    if (chapterDescription):
        paragraphWrite.style = document.styles['Chapter Description']
    else:
        if data.PublishOrEdit == 'toEdit':
            paragraphWrite.style = document.styles['Paragraph Editor']
            font = paragraphWrite.style.font
            font.size = Pt(12)
        else:
            paragraphWrite.style = document.styles['Paragraph']
            font = paragraphWrite.style.font
            font.size = Pt(11)

    paragraph_format = paragraphWrite.paragraph_format
    paragraph_format.keep_together = keep_together

    if indentFirstLine:
        paragraph_format.first_line_indent = Pt(18)

    typeNextParagraph = setParagraphType(paragrapghReadNext)

    if (typeNextParagraph == PART_TEXT or typeNextParagraph == CHAPTER_TEXT or typeNextParagraph == EXTRA_SECTION) :
        if previousTitle == EXTRA_SECTION:
            document.add_section()
            document.add_section()
        else:
            document.add_section()

def setSceneBreakFormat(document, paragraph):
    paragraphWrite = document.add_paragraph(paragraph)
    paragraphWrite.style = document.styles['Scene Break']

def setSectionFormat(document, dataFromHtml, doc):
    section = document.sections[0]

    #paragraphsCount = len(doc.paragraphs)
    #paragraphReadChapterCount = 0
    #countWords = 0
    #while paragraphReadChapterCount < paragraphsCount:
     #   paragraphRead = doc.paragraphs[paragraphReadChapterCount].text.strip()
      #  count = paragraphRead.split(" ")
       # countWords = countWords + len(count)
        #paragraphReadChapterCount = paragraphReadChapterCount + 1

    countWords = 100000
    if dataFromHtml.PageSize == '5" x 8" (12.7 x 20.32 cm)':
        section.page_width = Inches(5)
        section.page_height = Inches(8)
    elif dataFromHtml.PageSize == '5.25" x 8"(13.34 x 20.32 cm)':
        section.page_width = Inches(5.25)
        section.page_height = Inches(8)
    elif dataFromHtml.PageSize =='5.5" x 8.5"(13.97 x 21.59 cm)':
        section.page_width = Inches(5.5)
        section.page_height = Inches(8.5)
    elif dataFromHtml.PageSize == '6" x 9" (15.24 x 22.86 cm)':
        section.page_width = Inches(6)
        section.page_height = Inches(9)
    elif dataFromHtml.PageSize == '7" x 10" x (17.78 x 25.4 cm)':
        section.page_width = Inches(7)
        section.page_height = Inches(10)
    elif dataFromHtml.PageSize == '8" x 10" (20.32 x 25.4 cm)':
        section.page_width = Inches(8)
        section.page_height = Inches(10)
    elif dataFromHtml.PageSize == 'Letter 8.5" x 11" (21,59 x 27,94 cm)':
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    elif dataFromHtml.PageSize == 'A4 8.27" x 11.69" (21 x 29,7 cm)':
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    if dataFromHtml.PublishOrEdit == 'toEdit':
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    else:
        section.right_margin = Inches(0.3)

        pageCount = countWords/250

        if pageCount < 150:
            marginCalc = 0.375
        elif pageCount < 300:
            marginCalc = 0.5
        elif pageCount < 500:
            marginCalc = 0.625
        elif pageCount < 700:
            marginCalc = 0.75
        elif pageCount < 828:
            marginCalc = 0.875
        else:
            marginCalc = 0.875

        margin = 0.375


        if margin > marginCalc:
            section.left_margin = Inches(margin)
        else:
            section.left_margin = Inches(marginCalc)

    section.different_first_page_header_footer = True
    document.settings.odd_and_even_pages_header_footer = True

    header = section.header
    #footer = section.footer
    evenHeader = section.even_page_header
    evenFooter = section.even_page_footer

    header.is_linked_to_previous = False

    textHeader = header.paragraphs[0]
    textHeader.text = dataFromHtml.bookTitle.upper()

    textEvenHeader = evenHeader.paragraphs[0]
    textEvenHeader.text = dataFromHtml.author.upper()
