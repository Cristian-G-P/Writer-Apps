import docx
import re
from .auxFunctions import *
from .values import *
from flask import Blueprint, render_template, request, flash, redirect, url_for, send_file, flash
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from collections import OrderedDict


from datetime import datetime


## PONER LAS PLABARAS MAS REPETIDAS SIN LOS ARTICULOS   - PRONOMBRES - PREPOSICIONES
# SHOW DONT TELL
# PAssive voice: The road was made,
# Long Sentences

class RemoveSpecialCharsValue:
    cleanString = ''
    isWordAccepted = False


class Results:
    paragraphSizeDic = dict()
    paragraphTextDic = dict()
    numberOfWordsOrdered = []
    adverbsCount = dict()
    wordsInChapter = []
    wordsInBook = int
    filterWords = dict()
    glueWords = dict()
    redundantPhrases = dict()
    fillerPhrases = dict()
    nominalizationWords = dict()



def analyzeGrammar(file, dataFromHtml):
    results = Results()

    doc = docx.Document('website/temporal/'+ file.filename)
    wordsCountInChapter = 0
    wordsCountInBook = 0
    chapterCount = 0
    previousWord = ''
    adverbsCount = dict()
    paragraphReadChapterCount = 0

    paragraphWords = dict()
    wordsInChapter = dict()

    SetGrammarSectionFormat(doc, dataFromHtml)


    paragraphSizeDic = dict()

    countChapter = False

    while paragraphReadChapterCount < len(doc.paragraphs):
        paragraphRead = doc.paragraphs[paragraphReadChapterCount].text

        paragraphReadChapterCount = paragraphReadChapterCount + 1
        aux = paragraphRead.strip()

        keys = aux.split(" ")
        wordsCountInBook = wordsCountInBook + len(keys)


        if keys[0] =='CHAPTER':
            wordsCountInChapter = 0
            countChapter = True
            chapterName = ''
            i = 0
            while i < len(keys):
                chapterName = chapterName + ' ' + keys[i]
                i = i + 1
            wordsInChapter[chapterName] = 0

        results.paragraphSizeDic[paragraphReadChapterCount] = len(keys)

        if len(keys) > PARAGRAPH_NORMAL_SIZE:
            results.paragraphTextDic[paragraphReadChapterCount] = paragraphRead[0:40]

        paragraphReadSplit = paragraphRead.split()
        paragraphReadSplitUpper = []
        ##To put in order the dictionary of words in paragraphs
        str3 = []

        ##iterates every word in the current paragraph
        for i in paragraphReadSplit:
            paragraphReadSplitUpper.append(i.upper())

            key = ' '
        #################################################################################
        ##Remove special characters
            regex = re.compile('[@_!#$%^&*()<>?/\|}{~:,.]')
            #for i in range(0, len(paragraphReadSplitUpper)):

            if (regex.search(i) == None):
                isWordAccepted = True
                key = i.upper()
            else:
                removeSpecialCharsValue = removeSpecialChars(i.upper())
                key = removeSpecialCharsValue.cleanString

            #if not (removeSpecialCharsValue.isWordAccepted):
             #       continue

            #################################################################################
            ##Chapter counter
            #if key == 'CHAPTER':
            #    chapterCount = chapterCount + 1
            #    wordsCountInChapter = 0

            if countChapter:
                wordsCountInChapter = wordsCountInChapter + 1
                #wordsInChapter[chapterName] = wordsCountInChapter

            #################################################################################
            ##Word repetition counter
            if key in paragraphWords.keys():
                paragraphWords[key] = paragraphWords[key] + 1
            else:
                paragraphWords[key] = 1
            ################################################################################

            # GLUE WORDS
            if key in GLUE_WORDS:
                if key in results.glueWords.keys():
                    results.glueWords[key] = results.glueWords[key] + 1
                else:
                    results.glueWords[key] = 1

                    ################################################################################



                #################################################################################
                # FILTER WORDS
            if key in FILTER_WORDS:
                if key in results.filterWords.keys():
                    results.filterWords[key] = results.filterWords[key] + 1
                else:
                    results.filterWords[key] = 1

            ###################################################################################

            # Calculate number of adverbs
            position = (len(key) - 2)
            if key[position:] == 'LY':

                if key in adverbsCount.keys():
                    adverbsCount[key] = adverbsCount[key] + 1
                else:
                    adverbsCount[key] = 1

###################################################################################
            # REDUNDANT PHRASES
            redundantPhrase = previousWord + ' ' + key
            if redundantPhrase in REDUNDANT_PHRASES:
                if key in results.redundantPhrases.keys():
                    results.redundantPhrases[redundantPhrase] = results.redundantPhrases[redundantPhrase] + 1
                else:
                    results.redundantPhrases[redundantPhrase] = 1
###################################################################################
            # FILLER PHRASES
            fillerPhrase = previousWord + ' ' + key
            if fillerPhrase in FILLER_PHRASES:

                if key in results.fillerPhrases.keys():
                    results.fillerPhrases[fillerPhrase] = results.fillerPhrases[fillerPhrase] + 1
                else:
                    results.fillerPhrases[fillerPhrase] = 1
###################################################################################
            # NOMINALIZATION WORDS
            if key in NOMINALIZATION_WORDS:
                if key in results.nominalizationWords.keys():
                    results.nominalizationWords[key] = results.nominalizationWords[key] + 1
                else:
                    results.nominalizationWords[key] = 1
###################################################################################
            previousWord = key
        if countChapter:
            wordsInChapter[chapterName] = wordsInChapter[chapterName] + wordsCountInChapter
        wordsCountInChapter = 0
###################################################################################

    results.paragraphWords = paragraphWords
    results.numberOfWordsOrdered = str3
    results.adverbsCount = adverbsCount
    results.wordsInChapter = wordsInChapter
    results.wordsInBook = wordsCountInBook
    displayResults(results, dataFromHtml)


def removeSpecialChars(stringToRemoveChars):
    response = RemoveSpecialCharsValue()
    stringCorrected = stringToRemoveChars

    regex = re.compile('[@_!#$%^&*()<>?/\|}{~:,.’]')
    isWordAccepted = False

    if (regex.search(stringToRemoveChars) == None):
        q = 0

    else:
        aux1 = stringToRemoveChars.replace('.', '')
        aux2 = aux1.replace(',', '')
        aux3 = aux2.replace(';', '')
        aux4 = aux3.replace('’', '')
        aux5 = aux4.replace('‘', '')
        aux6 = aux5.replace(':', '')
        aux7 = aux6.replace('?', '')
        aux8 = aux7.replace('_', '')
        aux9 = aux8.replace('!', '')
        aux10 = aux9.replace('(', '')
        aux11 = aux10.replace(')', '')
        aux12 = aux11.replace('"', '')
        aux13 = aux12.replace('-', '')
        aux14 = aux13.replace('“', '')
        aux15 = aux14.replace('“', '')


    response.cleanString = aux15

    response.isWordAccepted = True

    return response


def displayResults(results, dataFromHtml):
    ###################################################################################
    for x in results.paragraphWords.keys():
        results.numberOfWordsOrdered.append(results.paragraphWords[x])

    #    results.numberOfWordsOrdered.sort(reverse=True);
    # list out keys and values separately
    key_list = list(results.paragraphWords.keys())
    val_list = list(results.paragraphWords.values())

    #######################################################################################
    ## print('#########OVER USED WORDS##########')
    for z in results.paragraphWords.keys():
        results.numberOfWordsOrdered.append(results.paragraphWords[z])

    # results.numberOfWordsOrdered.sort(reverse=True);
    # list out keys and values separately
    key_list = list(results.paragraphWords.keys())
    val_list = list(results.paragraphWords.values())

    val_list2 = list(results.paragraphWords.values())
    val_list2.sort(reverse=True)

    mostUsedWordsDic = dict()

    wordsCounter = 0
    # for i in results.numberOfWordsOrdered:
    ##Si dos palabras tiene la misma cantidad de repeticiones, muestra una sola
    for i in val_list2:
        position = val_list.index(i)
        mostUsedWordsDic[key_list[position]] = val_list[position]

        wordsCounter = wordsCounter + 1
        if wordsCounter > 250:
            ##ESTE VALOR DEBERIA SER CONSTANTE
            break

    results.mostUsedWordsDic = mostUsedWordsDic

    ##############################################################

    printDoc(results, dataFromHtml)


def printDoc(results, dataFromHtml):
    document = Document()
    section = document.sections[0]
    header = section.header
    footer = section.footer
    textHeader = header.paragraphs[0]
    textHeader.text = WEBSITE_ADDRESS + '\tGanymede Text Analyzer'
    textFooter = footer.paragraphs[0]
    textFooter.text = WEBSITE_ADDRESS

#########################################################################################

    setTitleFormatGrammar(document, dataFromHtml.bookTitle.upper())
    paragraphWrite = document.add_paragraph(GANYMEDE_DESCRIPTION_TEXT)
    paragraphWrite = document.add_paragraph('\t\t\t\t\t\t\t\t Ganymede Text Analyzer')



#########################################################################################

    paragraph = 'FILLER WORDS'
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)

    paragraph = FILLER_WORDS_TEXT
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 5'
    row = table.rows[0].cells
    row[0].text = 'FILLER WORDS'
    row[1].text = 'COUNT'

    for x in FILLER_WORDS:
        if x in results.paragraphWords.keys():
            row = table.add_row().cells
            row[0].text = str(x)
            row[1].text = str(results.paragraphWords[x])

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check these links for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://www.98thpercentile.com/blog/what-are-verbal-fillers-how-do-they-affect-your-speech/',
              '98thPercentile.com')
    printLink(document, 'https://crowwriter.com/cutting-filler-words-in-writing/', 'crowwriter.com')
    printLink(document, 'https://becomeawritertoday.com/filler-words-list/', 'becomeawritertoday.com')
    printLink(document, 'https://blog.wordvice.com/avoid-fillers-powerful-writing/', 'blog.wordadvice.com')
    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')

##################################################################
    paragraph = 'FILTER WORDS'
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)

    paragraph = FILTER_WORDS_TEXT
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 5'
    row = table.rows[0].cells
    row[0].text = 'FILTER WORDS'
    row[1].text = 'COUNT'

    for x in results.filterWords:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(x)
        row[1].text = str(results.filterWords[x])

        table.style = 'Medium Grid 3 Accent 5'

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check these links for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://michaeljmcdonagh.wordpress.com/2014/01/28/writer-unfiltered/', 'michaeljmcdonagh.wordpress.com')
    printLink(document, 'https://www.masterclass.com/articles/how-to-avoid-unnecessary-filter-words-in-your-writing', 'masterclass.com')
    printLink(document, 'https://www.rabbitwitharedpen.com/blog/filter-words-in-fiction','rabbitwitharedpen.com')
    printLink(document, 'https://heebel.com/2017/05/28/a-few-easy-strategies-to-remove-those-pesky-filter-words-that-fog-up-your-writing/', 'heebel.com')


##################################################################
    paragraph = 'TAG WORDS'
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = TAG_WORDS_TEXT
    paragraphWrite = document.add_paragraph(paragraph)
    paragraph = TAG_WORDS_TEXT_2
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 5'
    row = table.rows[0].cells
    row[0].text = 'TAG WORDS'
    row[1].text = 'COUNT'

    for x in TAG_WORDS:
        if x in results.paragraphWords.keys():
            paragraph = x

            row = table.add_row().cells
            row[0].text = str(x)
            row[1].text = str(results.paragraphWords[x])

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check these links for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://www.scribophile.com/academy/he-said-she-said-dialog-tags-and-using-them-effectively', 'scribophile.com')
    printLink(document, 'https://www.michellereneemiller.com/dialogue-tags/', 'michellereneemiller.com')
    printLink(document, 'https://www.nownovel.com/blog/dialogue-words-other-words-for-said/', 'nownovel.com')
    printLink(document, 'https://thewritepractice.com/dialogue-tags/', 'thewritepractice.com')
    printLink(document, 'https://www.wattpad.com/948375492-data-for-you-dear-tips-0-10', 'wattpad.com')


##################################################################
    paragraph = 'GLUE WORDS'
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)

    paragraph = GLUE_WORDS_TEXT
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 5'
    row = table.rows[0].cells
    row[0].text = 'GLUE WORDS'
    row[1].text = 'COUNT'

    for x in results.glueWords:
        row = table.add_row().cells
        row[0].text = str(x)
        row[1].text = str(results.glueWords[x])

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check this link for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://readable.com/blog/what-are-glue-words-and-how-do-they-affect-readability/',
              'readable.com')

##################################################################
    printWithDescription(document, 'NOMINALIZATION WORD', NOMINALIZATION_WORDS_TEXT, 'NOMINALIZATION WORDS', 'COUNT','DESCRIPTION', results.nominalizationWords, NOMINALIZATION_WORDS)

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check these links for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://www.thoughtco.com/nominalization-in-grammar-1691430', 'thoughtco.com')
    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')
##################################################################
    paragraph = 'ADVERBS ENDING IN LY'
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)

    paragraph = ADVERBS_TEXT
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)

    row = table.rows[0].cells
    row[0].text = 'ADVERBS (-LY)'
    row[1].text = 'COUNT'

    for x in results.adverbsCount:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(x)
        row[1].text = str(results.adverbsCount[x])

        table.style = 'Medium Grid 3 Accent 5'

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check these links for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://glcubel.com/2020/05/12/a-guide-to-removing-adverbs/', 'glcubel.com')
##################################################################
    paragraph = 'OVERUSED WORDS (PART I)'
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = OVERUSED_WORDS_TEXT
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)

    row = table.rows[0].cells
    row[0].text = 'WORD'
    row[1].text = 'COUNT'

    for x in results.mostUsedWordsDic:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(x)
        row[1].text = str(results.mostUsedWordsDic[x])

        table.style = 'Medium Grid 3 Accent 5'
##################################################################
    paragraph = 'OVERUSED WORDS (PART II)'
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = OVERUSED_WORDS_TEXT_2
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)

    row = table.rows[0].cells
    row[0].text = 'WORD'
    row[1].text = 'COUNT'

    for x in results.mostUsedWordsDic:
        if x not in COMMON_WORDS:
         # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(x)
            row[1].text = str(results.mostUsedWordsDic[x])

    table.style = 'Medium Grid 3 Accent 5'

##################################################################
    paragraph = ('UNIQUE WORDS ')
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = UNIQUE_WORDS_TEXT
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 5'

    row = table.rows[0].cells
    row[0].text = 'NUMBER OF UNIQUE WORDS'
    row[1].text = str(len(results.numberOfWordsOrdered))
##################################################################
    printWithDescription(document, 'FILLER PHRASES', FILLER_PHRASES_TEXT, 'FILLER PHRASE', 'COUNT', 'DESCRIPTION', results.fillerPhrases, FILLER_PHRASES)

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check this link for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')
##################################################################
    printWithDescription(document, 'REDUNDANT PHRASE', REDUNDANT_PHRASES_TEXT, 'REDUNDANT PHRASE', 'COUNT', 'DESCRIPTION', results.redundantPhrases, REDUNDANT_PHRASES)

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check this link for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')
##################################################################
    paragraph = ('PARAGRAPH SIZE')
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = PARAGRAPH_SIZE_TEXT
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=3)
    table.style = 'Medium Grid 3 Accent 5'

    row = table.rows[0].cells
    row[0].text = 'PARAGRAPH'
    row[1].text = 'WORD COUNT'
    row[2].text = 'STARTS WITH'

    for x in results.paragraphSizeDic:
        if (results.paragraphSizeDic[x] > PARAGRAPH_NORMAL_SIZE):
            row = table.add_row().cells

            row[0].text = str(x)
            row[1].text = str(results.paragraphSizeDic[x])
            row[2].text = str(results.paragraphTextDic[x])

    ##################################################################
    paragraph = ('CHAPTER SIZE')
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraphWrite = document.add_paragraph(CHAPTER_SIZE_TEXT)

    table = document.add_table(rows=1, cols=3)

    row = table.rows[0].cells
    row[0].text = 'CHAPTER'
    row[1].text = 'WORD COUNT'
    row[2].text = 'COMMENTS'

    for x in results.wordsInChapter:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(x)
        row[1].text = str(results.wordsInChapter[x])
        if (results.wordsInChapter[x] > CHAPTER_NORMAL_SIZE):
            row[2].text = 'Exceeded'
        else:
            row[2].text = 'OK'

        table.style = 'Medium Grid 3 Accent 5'
##################################################################
    paragraph = ('BOOK LENGTH')
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraphWrite = document.add_paragraph(BOOK_LENGTH_TEXT)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 5'
    row = table.rows[0].cells
    row[0].text = 'WORD COUNT'
    row[1].text = str(results.wordsInBook)
    paragraphWrite = document.add_paragraph('')

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 5'

    row = table.rows[0].cells
    row[0].text = 'BOOK TYPE'
    row[1].text = 'WORDS'

    row = table.add_row().cells
    row = table.rows[1].cells
    row[0].text = 'SHORT STORY'
    row[1].text = '20000 - 50000'

    row = table.add_row().cells
    row = table.rows[2].cells
    row[0].text = 'NOVELLA'
    row[1].text = '20000 - 50000'

    row = table.add_row().cells
    row = table.rows[3].cells
    row[0].text = 'NOVEL - PAPERBACK'
    row[1].text = '50000 - 100000'

    row = table.add_row().cells
    row = table.rows[4].cells
    row[0].text = 'NOVEL - HARDBACK'
    row[1].text = '50000 - 800000'
##################################################################
    printSectionWithoutCases(document,'CLUNKY CONSTRUCTION', CLUNKY_CONSTRUCTION_TEXT, 'CLUNKY CONSTRUCTION', 'DESCRIPTION', CLUNKY_CONSTRUCTIONS)

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check this link for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')

##################################################################
    printSectionWithoutCases(document,'EMPTY PHRASES', EMPTY_PHRASES_TEXT, 'EMPTY PHRASE', 'DESCRIPTION', EMPTY_PHRASES)

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check this link for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')
##################################################################
    printSectionWithoutCases(document, 'NEGATIVE CONSTRUCTIONS', NEGATIVE_CONSTRUCTION_TEXT, 'NEGATIVE CONSTRUCTION', 'DESCRIPTION', NEGATIVE_CONSTRUCTIONS)

    paragraphWrite = document.add_paragraph('')
    paragraph = ('Check this link for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://smartblogger.com/filler-words/', 'smartblogger.com')
##################################################################
    printSectionWithoutCases(document, 'FILLER PHRASES - PART II', FILLER_PHRASES_TEXT_2, 'FILLER PHRASE', 'DESCRIPTION', FILLER_PHRASES2)
 ##################################################################

    paragraph = ('SENTENCE LENGTH')
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = SENTENCE_LENGTH_TEXT
    paragraphWrite = document.add_paragraph(paragraph)

    paragraph = ('Check this link for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://www.thoughtco.com/sentence-length-grammar-and-composition-1691948', 'thoughtco.com')
##################################################################
    paragraph = ('PASSIVE VOICE')
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = PASSIVE_VOICE_TEXT
    paragraphWrite = document.add_paragraph(paragraph)
    font = paragraphWrite.style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.underline = False

    paragraph = ('Check this link for more information:')
    paragraphWrite = document.add_paragraph(paragraph)
    printLink(document, 'https://yoast.com/the-passive-voice-what-is-it-and-how-to-avoid-it/', 'yoast.com')

########################################################################################################
    saveGanymedeLog(dataFromHtml.bookTitle)

    document.save('website/temporal/' + dataFromHtml.bookTitle + ' - Analyzed by Ganymede.docx')

########################################################################################################
def printLink(document, link, text):

    p = document.add_paragraph(style='List Bullet')
    hyperlink = add_hyperlink(p, link, text, '4122ff', True)
    font = p.style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

def printSection(document, paragraph, text, col1, col2, results):

    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = text
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 5'

    row = table.rows[0].cells
    row[0].text = col1
    row[1].text = col2

    for x in results:
        row = table.add_row().cells
        row[0].text = str(x)
        row[1].text = str(results[x])

#######################################################################################
def printWithDescription(document, paragraph, text, col1, col2, col3, results, dic):
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = text
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=3)
    table.style = 'Medium Grid 3 Accent 5'

    row = table.rows[0].cells
    row[0].text = col1
    row[1].text = col2
    row[2].text = col3


    for x in results:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(x)
        row[1].text = str(results[x])
        row[2].text = str(dic[x])


def printSectionWithoutCases (document, title, text, col1, col2, dic):
    paragraph = title
    paragraphWrite = document.add_paragraph('')
    setTitleFormatGrammar(document, paragraph)
    paragraph = text
    paragraphWrite = document.add_paragraph(paragraph)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 3 Accent 5'

    row = table.rows[0].cells
    row[0].text = col1
    row[1].text = col2

    for x in dic:
        row = table.add_row().cells
        row[0].text = str(x)
        row[1].text = str(dic[x])

############################################################
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    #if not underline:
     # u = docx.oxml.shared.OxmlElement('w:u')
      #u.set(docx.oxml.shared.qn('w:val'), 'none')
      #rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

############################################################
def SetGrammarSectionFormat(document, dataFromHtml):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    header = section.header
    textHeader = header.paragraphs[0]
    textHeader.text = dataFromHtml.bookTitle + "LIBRO"

def saveGanymedeLog(book):
    now = datetime.now()
    date_time = now.strftime("%m/%d/%Y, %H:%M:%S")
    print("date and time:", date_time)

    doc = docx.Document('website/temporal/GanymedeLogs.docx')

    #fullText = []
    #for para in doc.paragraphs:
     #   fullText.append('crrrrrr' + para.text)
      #  print(fullText)

    paragraph = date_time + ' --- ' + book
    doc.add_paragraph(paragraph)
    doc.save('website/temporal/GanymedeLogs.docx')


